"""Resume parsing utilities for PDF and DOCX files.

Extracts plain text from uploaded resume files so the content can be
stored and used by the AI agent to inform job search recommendations.
"""

import io
import logging
import os
from pathlib import Path

logger = logging.getLogger(__name__)

# Maximum file size: 10 MB
MAX_FILE_SIZE = 10 * 1024 * 1024

ALLOWED_EXTENSIONS = {".pdf", ".docx"}


def allowed_file(filename: str) -> bool:
    """Check whether *filename* has an allowed extension."""
    ext = Path(filename).suffix.lower()
    return ext in ALLOWED_EXTENSIONS


def parse_resume(file_bytes: bytes, filename: str) -> str:
    """Parse a resume file and return its text content.

    Args:
        file_bytes: Raw bytes of the uploaded file.
        filename: Original filename (used to determine format).

    Returns:
        Extracted plain text from the resume.

    Raises:
        ValueError: If the file type is unsupported or the file is too large.
        RuntimeError: If parsing fails.
    """
    if len(file_bytes) > MAX_FILE_SIZE:
        raise ValueError(f"File too large ({len(file_bytes)} bytes). Maximum is {MAX_FILE_SIZE} bytes.")

    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _parse_pdf(file_bytes)
    elif ext == ".docx":
        return _parse_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Allowed types: {', '.join(ALLOWED_EXTENSIONS)}")


def _parse_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file using PyMuPDF."""
    try:
        import pymupdf
    except ImportError:
        raise RuntimeError("pymupdf is not installed. Run: uv add pymupdf")

    try:
        doc = pymupdf.open(stream=file_bytes, filetype="pdf")
        pages = []
        for page in doc:
            text = page.get_text()
            if text.strip():
                pages.append(text.strip())
        doc.close()

        if not pages:
            raise ValueError("Could not extract any text from the PDF. The file may be image-based or empty.")

        full_text = "\n\n".join(pages)
        logger.info("Parsed PDF: %d pages, %d chars", len(pages), len(full_text))
        return full_text
    except ValueError:
        raise
    except Exception as e:
        raise RuntimeError(f"Failed to parse PDF: {e}")


def _parse_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx is not installed. Run: uv add python-docx")

    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract text from tables (common in resumes)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        if not paragraphs:
            raise ValueError("Could not extract any text from the DOCX file. The file may be empty.")

        full_text = "\n".join(paragraphs)
        logger.info("Parsed DOCX: %d paragraphs, %d chars", len(paragraphs), len(full_text))
        return full_text
    except ValueError:
        raise
    except Exception as e:
        raise RuntimeError(f"Failed to parse DOCX: {e}")


def get_resume_dir() -> Path:
    """Return the directory where uploaded resumes are stored."""
    from backend.data_dir import get_data_dir
    resume_dir = get_data_dir() / "resumes"
    resume_dir.mkdir(parents=True, exist_ok=True)
    return resume_dir


def save_resume(file_bytes: bytes, filename: str) -> Path:
    """Save resume file to the data directory and return the path.

    Overwrites any existing file with the same name.
    """
    safe_name = Path(filename).name  # Strip any directory components
    dest = get_resume_dir() / safe_name
    dest.write_bytes(file_bytes)
    logger.info("Saved resume to %s (%d bytes)", dest, len(file_bytes))
    return dest


def get_saved_resume() -> dict | None:
    """Return info about the currently saved resume, or None if no resume exists.

    Returns:
        Dict with 'filename', 'path', and 'size' keys, or None.
    """
    resume_dir = get_resume_dir()
    files = sorted(resume_dir.iterdir(), key=lambda f: f.stat().st_mtime, reverse=True)
    resume_files = [f for f in files if f.suffix.lower() in ALLOWED_EXTENSIONS]
    if not resume_files:
        return None
    latest = resume_files[0]
    return {
        "filename": latest.name,
        "path": str(latest),
        "size": latest.stat().st_size,
    }


def get_resume_text() -> str | None:
    """Read and parse the most recently saved resume, returning its text.

    Returns None if no resume is saved.
    """
    info = get_saved_resume()
    if not info:
        return None
    path = Path(info["path"])
    file_bytes = path.read_bytes()
    return parse_resume(file_bytes, info["filename"])


def delete_resume() -> bool:
    """Delete all saved resume files. Returns True if any were deleted."""
    resume_dir = get_resume_dir()
    deleted = False
    for f in resume_dir.iterdir():
        if f.suffix.lower() in ALLOWED_EXTENSIONS:
            f.unlink()
            deleted = True
            logger.info("Deleted resume: %s", f)
    return deleted
